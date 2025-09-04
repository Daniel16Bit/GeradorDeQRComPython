import qrcode
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Image as RLImage, Spacer
from docx import Document
from docx.shared import Inches
from PIL import Image

# ====== CONFIGURAÇÕES ======
url = "https://drive.google.com/drive/folders/EXEMPLO_DO_LINK"
img_file = "qrcode_image.png"
pdf_file = "qrcode_document.pdf"
docx_file = "qrcode_document.docx"
qr_size = 10  # maior = mais resolução

# ====== GERAR QR CODE ======
qr = qrcode.QRCode(
    version=None,  # ajuste automático
    error_correction=qrcode.constants.ERROR_CORRECT_H,
    box_size=qr_size,
    border=4
)
qr.add_data(url)
qr.make(fit=True)

img = qr.make_image(fill_color="black", back_color="white").convert('RGB')
img.save(img_file)

# ====== GERAR PDF ======
pdf = SimpleDocTemplate(pdf_file, pagesize=A4)
pdf_width, pdf_height = A4
qr_width = 200

pdf_elements = [
    Spacer(1, 100),
    RLImage(img_file, width=qr_width, height=qr_width),
]

pdf.build(pdf_elements)

# ====== GERAR WORD ======
doc = Document()
doc.add_heading('QR Code Gerado', level=1)
doc.add_picture(img_file, width=Inches(2.5))
doc.add_paragraph("Link: " + url)
doc.save(docx_file)

print("QR Code gerado com sucesso:")
print(f"- Imagem: {img_file}")
print(f"- PDF: {pdf_file}")
print(f"- Word: {docx_file}")
